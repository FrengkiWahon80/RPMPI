import io
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- CONFIG HALAMAN ---
st.set_page_config(page_title="RPM DEEP LEARNING AI", page_icon="🧠", layout="wide")

# --- LOGIKA AI: GENERATOR KONTEN DETAIL ---
# Mengintegrasikan 3 Prinsip Pembelajaran & 3 Pengalaman Belajar ke dalam konten default
def generate_deep_learning_content(topik, subtopik):
    return {
        "awal": (
            f"1. **Orientasi (Kesadaran/Mindful)**: Guru mengajak peserta didik mengambil jeda hening sejenak untuk melatih napas sadar (mindful breathing), menyadari kehadiran diri, dilanjutkan dengan doa pembuka bersama.\n"
            f"2. **Apersepsi**: Guru mengaitkan konsep {topik} dengan dinamika atau pengalaman nyata yang dialami peserta didik sehari-hari.\n"
            f"3. **Motivasi (Bermakna/Meaningful)**: Guru menguraikan bagaimana nilai {subtopik} penting untuk menumbuhkan karakter mulia di dalam kehidupan bermasyarakat.\n"
            f"4. **Penyampaian Tujuan**: Menjelaskan kompetensi akhir yang akan dicapai dalam proses pembelajaran mendalam ini."
        ),
        "inti": {
            "surface": (
                f"**Pengalaman Belajar 1: Memahami (Pemerolehan/Surface)**:\n"
                f"- Peserta didik menyimak media literatur atau video interaktif secara terfokus mengenai {topik}.\n"
                f"- Mengidentifikasi definisi dasar, elemen penting, dan gagasan utama dari {subtopik}.\n"
                f"- Guru memberikan kuis interaktif singkat yang menyenangkan (Menggembirakan/Joyful) untuk memastikan pemahaman dasar terbentuk."
            ),
            "deep": (
                f"**Pengalaman Belajar 2: Mengaplikasi (Pengolahan/Deep)**:\n"
                f"- Peserta didik berkolaborasi dalam kelompok kecil untuk mengkritisi isu nyata atau studi kasus yang berhubungan dengan {subtopik}.\n"
                f"- Melakukan penalaran kritis dan diskusi mendalam untuk menyintesis solusi berbasis nilai-nilai moral/spiritual.\n"
                f"- Guru memfasilitasi sesi tanya jawab antar-kelompok yang dinamis dan bermakna."
            ),
            "transfer": (
                f"**Pengalaman Belajar 2: Mengaplikasi (Penerapan/Transfer)**:\n"
                f"- Peserta didik merancang karya inovatif berupa poster digital, infografis, atau rancangan aksi sosial mandiri yang mengejawantahkan nilai {topik}.\n"
                f"- Menyajikan karya tersebut kepada rekan sejawat untuk mendapatkan umpan balik konstruktif."
            )
        },
        "penutup": (
            f"1. **Pengalaman Belajar 3: Merefleksi**: Peserta didik menuliskan jurnal refleksi pribadi yang menjawab pertanyaan bermakna tentang perasaan, pemahaman baru, serta komitmen tindakan nyata mereka setelah mempelajari {subtopik}.\n"
            f"2. **Umpan Balik**: Guru mengapresiasi keaktifan peserta didik secara hangat dan memberikan penguatan moral.\n"
            f"3. **Doa Penutup**: Menutup proses pembelajaran dengan doa syukur atas hikmah pembelajaran yang didapatkan hari ini."
        ),
        "rubrik": (
            "1. **Pemahaman Konsep (Memahami - 30%)**: Ketepatan penjelasan fakta dasar dan landasan teori.\n"
            "2. **Aplikasi Praktis & Kreativitas (Mengaplikasi - 40%)**: Kualitas argumen, relevansi rancangan proyek, dan orisinalitas ide.\n"
            "3. **Kedalaman Refleksi (Merefleksi - 20%)**: Ketajaman analisis diri, kejujuran menuliskan jurnal, serta kejelasan komitmen pribadi.\n"
            "4. **Karakter & Kolaborasi (10%)**: Kerja sama, komunikasi yang santun, serta keterlibatan aktif selama proses."
        )
    }

# --- FUNGSI EKSPOR WORD (SESUAI TEMPLATE PDF) ---
def set_cell_background(cell, fill_color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tc_pr.append(shd)

def export_word_rpm(data):
    doc = Document()
    
    # Judul Utama
    title = doc.add_paragraph()
    run = title.add_run("PERENCANAAN PEMBELAJARAN MENDALAM (DEEP LEARNING)")
    run.font.bold = True
    run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tabel Identitas Atas
    tbl_top = doc.add_table(rows=5, cols=2)
    tbl_top.style = 'Table Grid'
    rows_top = [
        ("SEKOLAH", f": {data['sekolah']}"),
        ("NAMA GURU", f": {data['guru']}"),
        ("MATA PELAJARAN", f": {data['mapel']}"),
        ("KELAS / SEMESTER", f": {data['kelas']}"),
        ("ALOKASI WAKTU", f": {data['durasi']}")
    ]
    for i, (k, v) in enumerate(rows_top):
        tbl_top.rows[i].cells[0].text = k
        tbl_top.rows[i].cells[1].text = v

    doc.add_paragraph() # Spacer

    # Tabel Utama Identifikasi, Desain, Pengalaman (Kini menggunakan 16 baris)
    tbl_main = doc.add_table(rows=16, cols=3)
    tbl_main.style = 'Table Grid'
    
    def fill_row(idx, col1, col2, col3):
        tbl_main.rows[idx].cells[0].text = col1
        tbl_main.rows[idx].cells[1].text = col2
        tbl_main.rows[idx].cells[2].text = col3
        set_cell_background(tbl_main.rows[idx].cells[1], "F2F2F2")

    # Bagian IDENTIFIKASI (8 Dimensi diintegrasikan pada baris Profil Lulusan)
    fill_row(0, "IDENTIFIKASI", "Peserta Didik", data['karakteristik'])
    fill_row(1, "", "Materi Pelajaran", f"Topik Utama: {data['topik']}\nSub-topik: {data['subtopik']}")
    fill_row(2, "", "Dimensi Profil Lulusan", data['profil'])
    fill_row(3, "", "Capaian Pembelajaran", data['cp'])
    fill_row(4, "", "Lintas Disiplin Ilmu", data['lintas'])
    fill_row(5, "", "Tujuan Pembelajaran", data['tujuan'])

    # Bagian DESAIN PEMBELAJARAN (Menyisipkan 3 Prinsip Pembelajaran & 4 Kerangka Pembelajaran)
    fill_row(6, "DESAIN PEMBELAJARAN", "Prinsip Pembelajaran (3 Prinsip)", data['prinsip'])
    fill_row(7, "", "Praktik Pedagogik (Kerangka 1)", data['pedagogis'])
    fill_row(8, "", "Kemitraan Pembelajaran (Kerangka 2)", data['kemitraan'])
    fill_row(9, "", "Lingkungan Pembelajaran (Kerangka 3)", data['lingkungan'])
    fill_row(10, "", "Pemanfaatan Digital (Kerangka 4)", data['digital'])

    # Bagian PENGALAMAN BELAJAR (Menerapkan 3 Pengalaman Belajar)
    fill_row(11, "PENGALAMAN BELAJAR", "Memahami (Awal & Surface)", f"{data['p_awal']}\n\n{data['p_inti_surf']}")
    fill_row(12, "", "Mengaplikasi (Deep & Transfer)", f"{data['p_inti_deep']}\n\n{data['p_inti_tran']}")
    fill_row(13, "", "Merefleksi (Penutup)", data['p_penutup'])

    # Bagian ASESMEN
    fill_row(14, "ASESMEN", "Teknik & Instrumen", data['as_teknik'])
    fill_row(15, "", "Rubrik Penilaian", data['as_rubrik'])

    # Tanda Tangan
    doc.add_paragraph("\n")
    ttd = doc.add_table(rows=3, cols=2)
    ttd.rows[0].cells[0].text = "Mengetahui,\nKepala Sekolah"
    ttd.rows[0].cells[1].text = f"Merdeka, ................ 2026\nGuru Mata Pelajaran"
    ttd.rows[2].cells[0].text = f"{data['kepsek']}\nNIP. {data['nip_kepsek']}"
    ttd.rows[2].cells[1].text = f"{data['guru']}\nNIP. {data['nip_guru']}"

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- INTERFACE UTAMA STREAMLIT ---
st.title("📄 RPM DEEP LEARNING GENERATOR")
st.markdown("Sesuai Template SMPN Tujuh Maret Hadakewa dengan Formula Pembelajaran Mendalam")

with st.sidebar:
    st.header("📋 Administrasi")
    sekolah = st.text_input("Sekolah", "SMPN Tujuh Maret Hadakewa")
    guru = st.text_input("Nama Guru", "Daniel Florensius Lako Wahon, S.S")
    nip_guru = st.text_input("NIP Guru", "19801032024211002")
    kepsek = st.text_input("Kepala Sekolah", "Fransiskus Bernardus Kedang Kaona, S.Fl")
    nip_kepsek = st.text_input("NIP Kepala Sekolah", "19800132006041015")
    mapel = st.text_input("Mata Pelajaran", "Agama Katolik dan Budi Pekerti")
    kelas = st.text_input("Kelas / Semester", "IX / 1")
    durasi = st.text_input("Alokasi Waktu", "2 JP (2 x 40 Menit)")

st.subheader("🔍 1. Identifikasi & Desain (Formula 8-3-48)")
c1, c2 = st.columns(2)
with c1:
    topik = st.text_input("Topik Utama", "Gereja yang Beriman")
    subtopik = st.text_input("Sub-topik", "Peran Serta dalam Hirarki")
    
    # Integrasi 8 Dimensi Profil Lulusan
    profil = st.multiselect("Dimensi Profil Lulusan (8 Dimensi)", 
                           ["Beriman & bertakwa", "Kewargaan", "Penalaran kritis", "Kreativitas", "Kemandirian", "Kolaborasi", "Komunikasi", "Kesehatan"],
                           default=["Beriman & bertakwa", "Penalaran kritis", "Kolaborasi"])
    
    # Integrasi 3 Prinsip Pembelajaran
    prinsip = st.multiselect("Prinsip Pembelajaran (3 Prinsip)", 
                            ["Kesadaran (Mindful)", "Bermakna (Meaningful)", "Menggembirakan (Joyful)"],
                            default=["Kesadaran (Mindful)", "Bermakna (Meaningful)", "Menggembirakan (Joyful)"])
    
    lintas = st.text_input("Lintas Disiplin Ilmu", "PPKn (Struktur Organisasi), Bahasa Indonesia (Literasi Kitab Suci)")
with c2:
    st.markdown("**4 Kerangka Pembelajaran:**")
    pedagogis = st.selectbox("1. Praktik Pedagogik", ["Pembelajaran Berbasis Masalah", "Pembelajaran Berbasis Proyek", "Inkuiri", "Kontekstual"])
    lingkungan = st.text_input("2. Lingkungan Pembelajaran", "Ruang kelas kondusif, inklusif, serta terkoneksi internet")
    digital = st.text_input("3. Pemanfaatan Digital", "Video Pembelajaran, Canva, Google Form untuk Refleksi")
    kemitraan = st.text_area("4. Kemitraan Pembelajaran", "Orang tua (diskusi iman), Tokoh Agama/Katekis (Narasumber)", height=68)

st.divider()
st.subheader("🚀 2. Pengalaman Belajar (Detail 3 Pengalaman)")

ai_content = generate_deep_learning_content(topik, subtopik)

# Mengatur tab sesuai dengan 3 Pengalaman Belajar (Memahami, Mengaplikasi, Merefleksi)
tab_paham, tab_aplikasi, tab_refleksi = st.tabs([
    "Memahami (Awal & Surface)", 
    "Mengaplikasi (Deep & Transfer)", 
    "Merefleksi (Penutup)"
])

with tab_paham:
    p_awal = st.text_area("Aktivitas Pembuka (Mindful & Relevan)", ai_content['awal'], height=150)
    p_surf = st.text_area("Pemerolehan Konsep (Surface)", ai_content['inti']['surface'], height=180)

with tab_aplikasi:
    st.info("Proses pengolahan materi secara mendalam hingga penerapan gagasan baru.")
    col_d, col_t = st.columns(2)
    p_deep = col_d.text_area("Pengolahan Ide (Deep)", ai_content['inti']['deep'], height=250)
    p_tran = col_t.text_area("Penerapan Aksi (Transfer)", ai_content['inti']['transfer'], height=250)

with tab_refleksi:
    p_penutup = st.text_area("Aktivitas Penutup (Refleksi & Komitmen)", ai_content['penutup'], height=180)

st.divider()
st.subheader("📊 3. Asesmen & Rubrik")
ca, cb = st.columns(2)
with ca:
    as_teknik = st.text_area("Teknik & Instrumen Penilaian", 
                            "1. Penilaian Diri/Jurnal Reflektif (As Learning)\n2. Observasi Kolaborasi (For Learning)\n3. Portofolio/Karya Kreatif (Of Learning)")
with cb:
    as_rubrik = st.text_area("Rubrik Penilaian Detail", ai_content['rubrik'], height=150)

# Pengumpulan data untuk diekspor ke format dokumen (.docx)
data_rpm = {
    "sekolah": sekolah, "guru": guru, "nip_guru": nip_guru, "kepsek": kepsek, "nip_kepsek": nip_kepsek,
    "mapel": mapel, "kelas": kelas, "durasi": durasi,
    "karakteristik": "Peserta didik memiliki latar belakang iman beragam serta memerlukan stimulasi aktif agar dapat berkolaborasi dan bernalar kritis.",
    "topik": topik, "subtopik": subtopik, 
    "profil": ", ".join(profil),
    "prinsip": ", ".join(prinsip),
    "cp": "Peserta didik memahami peran, tugas, dan tanggung jawabnya secara aktif dalam kehidupan bersama.",
    "lintas": lintas, "tujuan": f"Peserta didik mampu memahami nilai {subtopik}, mengaplikasikannya dalam proyek, serta merefleksikan nilai tersebut dalam kehidupan sehari-hari.",
    "pedagogis": pedagogis, 
    "kemitraan": kemitran, 
    "lingkungan": lingkungan,
    "digital": digital, 
    "p_awal": p_awal, 
    "p_inti_surf": p_surf, 
    "p_inti_deep": p_deep, 
    "p_inti_tran": p_tran,
    "p_penutup": p_penutup, 
    "as_teknik": as_teknik, 
    "as_rubrik": as_rubrik
}

if st.button("📝 Generate & Download RPM Mendalam (.docx)", type="primary"):
    file_word = export_word_rpm(data_rpm)
    st.download_button(
        label="📥 Unduh Sekarang",
        data=file_word,
        file_name=f"RPM_Mendalam_{topik.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
